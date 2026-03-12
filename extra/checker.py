
import os
import PyPDF2
import docx
import pandas as pd

def extract_pdf_stats(path):
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return f"{len(reader.pages)} pages"
    except Exception as e:
        return str(e)

def extract_docx_details(path):
    try:
        doc = docx.Document(path)
        return f"{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables"
    except Exception as e:
        return str(e)

files = [
    r"c:\Users\user\Desktop\folders\mega\extra\Electrical_equipments.docx",
    r"c:\Users\user\Desktop\folders\mega\extra\OHL Substn.pdf"
]

for f in files:
    print(f"File: {os.path.basename(f)}")
    if f.endswith(".pdf"):
        print(f"  Stats: {extract_pdf_stats(f)}")
    elif f.endswith(".docx"):
        print(f"  Stats: {extract_docx_details(f)}")
