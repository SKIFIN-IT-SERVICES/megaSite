
import os
import PyPDF2
import docx
import pandas as pd

def extract_pdf(path):
    text = ""
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        text = f"Error reading PDF {path}: {e}"
    return text

def extract_docx(path):
    text = ""
    try:
        doc = docx.Document(path)
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Extract tables
        for table in doc.tables:
            text += "\n[TABLE START]\n"
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                text += row_text + "\n"
            text += "[TABLE END]\n"
    except Exception as e:
        text = f"Error reading DOCX {path}: {e}"
    return text

def extract_xlsx(path):
    text = ""
    try:
        df = pd.read_excel(path)
        text = df.to_string()
    except Exception as e:
        text = f"Error reading XLSX {path}: {e}"
    return text

files = [
    r"c:\Users\user\Desktop\folders\mega\extra\Elec and Inst Profile_260309_125407.pdf",
    r"c:\Users\user\Desktop\folders\mega\extra\Electical_Instr.xlsx",
    r"c:\Users\user\Desktop\folders\mega\extra\Electrical_equipments.docx",
    r"c:\Users\user\Desktop\folders\mega\extra\Hospitiality.docx",
    r"c:\Users\user\Desktop\folders\mega\extra\OHL Substn.pdf"
]

output_file = r"c:\Users\user\Desktop\folders\mega\extra\extracted_content.txt"

with open(output_file, "w", encoding="utf-8") as out:
    for f in files:
        out.write(f"--- CONTENT OF {os.path.basename(f)} ---\n")
        if f.endswith(".pdf"):
            out.write(extract_pdf(f))
        elif f.endswith(".docx"):
            out.write(extract_docx(f))
        elif f.endswith(".xlsx"):
            out.write(extract_xlsx(f))
        out.write("\n\n")

print(f"Content extracted to {output_file}")
